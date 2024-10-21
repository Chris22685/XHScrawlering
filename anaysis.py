import pandas as pd
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
import os

# 整理正文内容为Word文档
def combine_excel_to_word(excel_path, output_path):
    """
    读取Excel文件中的'正文'列内容，将其合并后写入Word文档，并设置文本格式
    
    Parameters:
    excel_path (str): Excel文件的路径
    output_path (str): 输出Word文件的路径
    """
    try:
        # 读取Excel文件
        df = pd.read_excel(excel_path)
        
        # 检查是否存在'正文'列
        if '正文' not in df.columns:
            raise ValueError("Excel文件中未找到'正文'列")
            
        # 合并所有正文内容，使用两个换行符分隔
        combined_text = '\n\n'.join(df['正文'].fillna('').astype(str))
        
        # 创建Word文档
        doc = Document()
        
        # 设置整个文档的基本格式
        # 设置默认字体
        doc.styles['Normal'].font.name = '宋体'
        doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        
        # # 添加标题
        # heading = doc.add_heading('合并的正文内容', level=1)
        # heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # heading_format = heading.runs[0]
        # heading_format.font.name = '黑体'
        # heading_format.font.size = Pt(16)
        # heading_format.font.bold = True
        # heading_format._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        
        # 添加正文内容并设置格式
        paragraph = doc.add_paragraph()
        paragraph_format = paragraph.paragraph_format
        
        # 设置段落格式
        paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE  # 1.5倍行距
        paragraph_format.space_after = Pt(0)  # 段后间距
        paragraph_format.first_line_indent = Inches(0.5)  # 首行缩进
        
        # 添加文本并设置字体
        run = paragraph.add_run(combined_text)
        font = run.font
        font.name = '宋体'
        font.size = Pt(12)
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        
        # 设置页面边距
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1.25)
            section.right_margin = Inches(1.25)
        
        # 保存文档
        doc.save(output_path)
        
        print(f"文件已成功保存到: {output_path}")
        
    except Exception as e:
        print(f"发生错误: {str(e)}")

# 设置文件路径
excel_path = '/Users/cris/Documents/Project/Find_Coder/actual/XHScrawler/result.xlsx'
output_path = 'combined_content.docx'

# 执行转换
combine_excel_to_word(excel_path, output_path)